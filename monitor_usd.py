import logging
import os
from datetime import datetime
from urllib.parse import urlparse
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import NoSuchElementException, TimeoutException
from docx import Document
from docx.shared import Pt, Inches
import tkinter as tk
from tkinter import font as tkFont
from docx2pdf import convert

# Logger configuration
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CustomDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Nome do Usuário")
        self.geometry("320x140")
        self.custom_font = tkFont.Font(family="Arial", size=12)
        self.bold_font = tkFont.Font(family="Arial", size=14, weight="bold")

        self.label = tk.Label(self, text="Por favor, digite seu nome:", font=self.bold_font)
        self.label.pack(padx=10, pady=10)

        self.entry = tk.Entry(self, font=self.custom_font, width=40)
        self.entry.pack(padx=10)

        self.ok_button = tk.Button(self, text="Iniciar", command=self.on_ok, font=self.custom_font)
        self.ok_button.pack(expand=True)

        self.result = None

    def on_ok(self):
        self.result = self.entry.get()
        self.destroy()

def ask_name():
    root = tk.Tk()
    root.withdraw()

    dialog = CustomDialog(root)
    root.wait_window(dialog)
    return dialog.result

def start_driver():
    chrome_options = Options()
    arguments = [
        '--lang=pt_BR',
        '--window-size=1366x768',
        '--incognito',
        '--ignore-certificate-errors',
        '--disable-gpu',
        '--disable-dev-shm-usage',
        '--ignore-ssl-errors',
        '--log-level=3',
    ]

    for argument in arguments:
        chrome_options.add_argument(argument)

    chrome_options.add_experimental_option('prefs', {
        'download.prompt_for_download': False,
        'profile.default_content_setting_values.notifications': 2,
        'profile.default_content_setting_values.automatic_downloads': 1,
    })

    driver = webdriver.Chrome(options=chrome_options)

    # Maximize the browser window
    driver.maximize_window()

    wait = WebDriverWait(driver, 10, poll_frequency=1, ignored_exceptions=[NoSuchElementException, TimeoutException])
    url = 'https://economia.uol.com.br/cotacoes/cambio/'

    return driver, wait, url

def get_exchange_rate(driver, url, wait):
    driver.get(url)

    try:
        close_button = wait.until(EC.visibility_of_element_located((By.CSS_SELECTOR, 'i.fc-close-icon')))
        driver.execute_script("arguments[0].click();", close_button)
        logger.info("Popup fechado com sucesso.")
    except TimeoutException:
        logger.error("Popup não apareceu ou não pôde ser fechado.")
    except NoSuchElementException:
        logger.error("Botão de fechar popup não encontrado.")
    except Exception as e:
        logger.error(f"Ocorreu um erro ao fechar o popup: {e}")

    try:
        exchange_rate_element = wait.until(EC.presence_of_element_located((By.XPATH,
                                                                           '/html/body/div[3]/article/div[2]/div/div[1]/div/div/div[1]/div[3]/div/div/div/div[2]/div[1]/div[2]/div/span[2]')))
        exchange_rate = exchange_rate_element.text.strip()
        exchange_rate = exchange_rate[:4]
        logger.info(f"Cotação do dia: {exchange_rate}")

        screenshot_path = 'screenshot.png'
        driver.get_screenshot_as_file(screenshot_path)
        if os.path.exists(screenshot_path):
            logger.info(f"Captura de tela realizada e salva em {screenshot_path}.")
        else:
            logger.error("Falha ao salvar a captura de tela.")

        return exchange_rate
    except TimeoutException:
        logger.error("Elemento com a cotação do dia não encontrado a tempo.")
        return None
    except NoSuchElementException:
        logger.error("Elemento com a cotação do dia não encontrado.")
        return None
    finally:
        driver.quit()

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            logger.info(f"Documento convertido para PDF como {pdf_path}.")
        else:
            logger.error(f"Falha ao salvar o PDF como {pdf_path}.")
    except Exception as e:
        logger.error(f"Erro ao converter o documento para PDF: {e}")

def create_exchange_rate_document(exchange_rate, url, user_name):
    today_date = datetime.now().strftime('%d-%m-%Y')
    url_hostname = urlparse(url).hostname

    title_text = f"Cotação do Dólar Comercial – R$ {exchange_rate} ({today_date})"
    body_text = f"""
O dólar está no valor de R$ {exchange_rate}, na data {today_date}.
Valor cotado no site {url_hostname}
Print da cotação atual
    """

    image_path = 'screenshot.png'

    doc = Document()

    def apply_arial_font(run, size, bold=False):
        run.font.name = 'Arial'
        run.font.size = size
        run.font.bold = bold

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title_text)
    apply_arial_font(title_run, Pt(14), bold=True)
    title_paragraph.alignment = 1

    doc.add_paragraph()

    body_paragraph = doc.add_paragraph()
    body_run = body_paragraph.add_run(body_text.strip())
    apply_arial_font(body_run, Pt(12))

    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.3))
    else:
        logger.warning("Screenshot não encontrada.")

    footer_paragraph = doc.add_paragraph()
    footer_run = footer_paragraph.add_run(f"Cotação feita por – {user_name}")
    apply_arial_font(footer_run, Pt(12))

    word_file_name = f'Cotação_Dólar_USD_{today_date}.docx'
    try:
        doc.save(word_file_name)
        logger.info(f"Criando Documento {word_file_name}.")
    except Exception as e:
        logger.error(f"Erro ao salvar o documento: {e}")
        return

    pdf_file_name = f'Cotação_Dólar_USD_{today_date}.pdf'
    try:
        convert_docx_to_pdf(word_file_name, pdf_file_name)
        logger.info(f"PDF gerado com sucesso: {pdf_file_name}")
        os.remove(image_path)
        logger.info("Favor verificar documento salvo na pasta.")
    except Exception as e:
        logger.error(f"Erro ao converter o documento para PDF: {e}")

def execute_tasks(user_name):
    logger.info(f"Olá, {user_name}! Vamos verificar a cotação do dólar e salvar em um arquivo em PDF.")
    driver, wait, url = start_driver()
    exchange_rate = get_exchange_rate(driver, url, wait)

    if exchange_rate:
        create_exchange_rate_document(exchange_rate, url, user_name)
    else:
        logger.error("Não foi possível obter a cotação do dólar.")

def main():
    user_name = ask_name()
    if user_name:
        execute_tasks(user_name)
        tk.Tk().quit()  # Close the root Tk instance

if __name__ == '__main__':
    main()